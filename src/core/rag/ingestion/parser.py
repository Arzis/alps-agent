"""文档解析器模块

将不同格式的文件 (PDF/DOCX/MD/TXT) 解析为 LlamaIndex Document 列表。
"""

from pathlib import Path

import structlog
from llama_index.core import Document

logger = structlog.get_logger()


class DocumentParser:
    """
    文档解析器 - 将不同格式的文件解析为纯文本

    支持的格式:
    - PDF: 使用 pypdf 解析
    - DOCX: 使用 python-docx 解析
    - Markdown: 直接读取
    - 纯文本: 直接读取
    """

    async def parse(self, file_path: str, file_type: str) -> list[Document]:
        """解析文档，返回 LlamaIndex Document 列表

        Args:
            file_path: 文件路径
            file_type: 文件类型 (如 ".pdf")

        Returns:
            list[Document]: 解析后的 Document 列表

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件类型
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 根据文件类型选择解析器
        parser_map = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
        }

        parser = parser_map.get(file_type.lower())
        if not parser:
            raise ValueError(f"不支持的文件类型: {file_type}")

        documents = await parser(path)

        logger.info(
            "document_parsed",
            file_path=file_path,
            file_type=file_type,
            num_pages=len(documents),
        )

        return documents

    async def _parse_pdf(self, path: Path) -> list[Document]:
        """解析 PDF 文件

        Args:
            path: PDF 文件路径

        Returns:
            list[Document]: 每页一个 Document
        """
        from pypdf import PdfReader
        import asyncio

        def _read_pdf():
            """同步读取 PDF"""
            reader = PdfReader(str(path))
            documents = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and text.strip():
                    documents.append(
                        Document(
                            text=text.strip(),
                            metadata={
                                "page_number": i + 1,  # 页码 (从1开始)
                                "total_pages": len(reader.pages),  # 总页数
                                "source": path.name,  # 文件名
                            },
                        )
                    )
            return documents

        # 在线程池中执行，避免阻塞事件循环
        return await asyncio.get_event_loop().run_in_executor(None, _read_pdf)

    async def _parse_docx(self, path: Path) -> list[Document]:
        """解析 Word 文档

        Args:
            path: DOCX 文件路径

        Returns:
            list[Document]: 整个文档一个 Document
        """
        from docx import Document as DocxDocument
        import asyncio

        def _read_docx():
            """同步读取 DOCX"""
            doc = DocxDocument(str(path))
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            return [
                Document(
                    text="\n\n".join(full_text),
                    metadata={"source": path.name},
                )
            ]

        return await asyncio.get_event_loop().run_in_executor(None, _read_docx)

    async def _parse_markdown(self, path: Path) -> list[Document]:
        """解析 Markdown 文件

        Args:
            path: MD 文件路径

        Returns:
            list[Document]: 整个文档一个 Document
        """
        content = path.read_text(encoding="utf-8")
        return [
            Document(
                text=content,
                metadata={"source": path.name},
            )
        ]

    async def _parse_text(self, path: Path) -> list[Document]:
        """解析纯文本文件

        Args:
            path: TXT 文件路径

        Returns:
            list[Document]: 整个文档一个 Document
        """
        content = path.read_text(encoding="utf-8")
        return [
            Document(
                text=content,
                metadata={"source": path.name},
            )
        ]
